"""
Make 'span' in tags dict a stack
maybe do the same for all tags in case of unclosed tags?
optionally use bs4 to clean up invalid html?

the idea is that there is a method that converts html files into docx
but also have api methods that let user have more control e.g. so they
can nest calls to something like 'convert_chunk' in loops

user can pass existing document object as arg 
(if they want to manage rest of document themselves)

How to deal with block level style applied over table elements? e.g. text align
"""
import re, argparse
import io, os
import urllib.request
from urllib.parse import urlparse
from html.parser import HTMLParser

import docx, docx.table
from docx import Document
from docx.shared import RGBColor, Pt, Inches
from docx.enum.text import WD_COLOR, WD_ALIGN_PARAGRAPH

from bs4 import BeautifulSoup

# values in inches
INDENT = 0.25
LIST_INDENT = 0.5
MAX_INDENT = 5.5 # To stop indents going off the page

def get_filename_from_url(url):
    return os.path.basename(urlparse(url).path)

def is_url(url):
    """
    Not to be used for actually validating a url, but in our use case we only 
    care if it's a url or a file path, and they're pretty distinguishable
    """
    parts = urlparse(url)
    return all([parts.scheme, parts.netloc, parts.path])

def fetch_image(url):
    """
    Attempts to fetch an image from a url. 
    If successful returns a bytes object, else returns None

    :return:
    """
    try:
        with urllib.request.urlopen(url) as response:
            # security flaw?
            return io.BytesIO(response.read())
    except urllib.error.URLError:
        return None

def remove_last_occurence(ls, x):
    ls.pop(len(ls) - ls[::-1].index(x) - 1)

def remove_whitespace(string):
    string = re.sub(r'\s*\n\s*', ' ', string)
    return re.sub(r'>\s{2+}<', '><', string)

def delete_paragraph(paragraph):
    # https://github.com/python-openxml/python-docx/issues/33#issuecomment-77661907
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

fonts = {
    'b': 'bold',
    'strong': 'bold',
    'em': 'italic',
    'i': 'italic',
    'u': 'underline',
    's': 'strike',
    'sup': 'superscript',
    'sub': 'subscript',
    'th': 'bold',
}

class HtmlToDocx(HTMLParser):

    def __init__(self):
        super().__init__()
        self.options = {
            'fix-html': True,
            'images': True,
            'tables': True,
            'styles': True,
        }
        self.table_row_selectors = [
            'table > tr',
            'table > thead > tr',
            'table > tbody > tr',
            'table > tfoot > tr'
        ]

    def set_initial_attrs(self, document=None):
        self.tags = {
            'span': [],
            'list': [],
        }
        if document:
            self.doc = document
        else:
            self.doc = Document()
        self.bs = self.options['fix-html'] # whether or not to clean with BeautifulSoup
        self.document = self.doc
        self.include_tables = True #TODO add this option back in?
        self.include_images = self.options['images']
        self.include_styles = self.options['styles']
        self.paragraph = None
        self.skip = False
        self.skip_tag = None
        self.instances_to_skip = 0

    def get_cell_html(self, soup):
        # Returns string of td element with opening and closing <td> tags removed
        # Cannot use find_all as it only finds element tags and does not find text which
        # is not inside an element
        return ' '.join([str(i) for i in soup.contents])

    def add_styles_to_paragraph(self, style):
        if 'text-align' in style:
            align = style['text-align']
            if align == 'center':
                self.paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif align == 'right':
                self.paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif align == 'justify':
                self.paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if 'margin-left' in style:
            margin = style['margin-left']
            units = re.sub(r'[0-9]+', '', margin)
            margin = int(float(re.sub(r'[a-z]+', '', margin)))
            if units == 'px':
                self.paragraph.paragraph_format.left_indent = Inches(min(margin // 10 * INDENT, MAX_INDENT))
            # TODO handle non px units

    def add_styles_to_run(self, style):
        if 'color' in style:
            if 'rgb' in style['color']:
                color = re.sub(r'[a-z()]+', '', style['color'])
                colors = [int(x) for x in color.split(',')]
            elif '#' in style['color']:
                color = style['color'].lstrip('#')
                colors = tuple(int(color[i:i+2], 16) for i in (0, 2, 4))
            else:
                colors = [0, 0, 0]
                # TODO map colors to named colors (and extended colors...)
                # For now set color to black to prevent crashing
            self.run.font.color.rgb = RGBColor(*colors)
            
        if 'background-color' in style:
            if 'rgb' in style['background-color']:
                color = color = re.sub(r'[a-z()]+', '', style['background-color'])
                colors = [int(x) for x in color.split(',')]
            elif '#' in style['background-color']:
                color = style['background-color'].lstrip('#')
                colors = tuple(int(color[i:i+2], 16) for i in (0, 2, 4))
            else:
                colors = [0, 0, 0]
                # TODO map colors to named colors (and extended colors...)
                # For now set color to black to prevent crashing
            self.run.font.highlight_color = WD_COLOR.GRAY_25 #TODO: map colors

    def parse_dict_string(self, string, separator=';'):
        new_string = string.replace(" ", '').split(separator)
        string_dict = dict([x.split(':') for x in new_string if ':' in x])
        return string_dict

    def handle_li(self):
        # check list stack to determine style and depth
        list_depth = len(self.tags['list'])
        if list_depth:
            list_type = self.tags['list'][-1]
        else:
            list_type = 'ul' # assign unordered if no tag

        if list_type == 'ol':
            list_style = "List Number"
        else:
            list_style = 'List Bullet'

        self.paragraph = self.doc.add_paragraph(style=list_style)            
        self.paragraph.paragraph_format.left_indent = Inches(min(list_depth * LIST_INDENT, MAX_INDENT))
        self.paragraph.paragraph_format.line_spacing = 1

    def add_image_to_cell(self, cell, image):
        # python-docx doesn't have method yet for adding images to table cells. For now we use this
        paragraph = cell.add_paragraph()
        run = paragraph.add_run()
        run.add_picture(image)

    def handle_img(self, current_attrs):
        if not self.include_images:
            self.skip = True
            self.skip_tag = 'img'
            return
        src = current_attrs['src']
        # fetch image
        src_is_url = is_url(src)
        if src_is_url:
            try:
                image = fetch_image(src)
            except urllib.error.URLError:
                image = None
        else:
            image = src
        # add image to doc
        if image:
            try:
                if isinstance(self.doc, docx.document.Document):
                    self.doc.add_picture(image)
                else:
                    self.add_image_to_cell(self.doc, image)
            except FileNotFoundError:
                image = None
        if not image:
            if src_is_url:
                self.doc.add_paragraph("<image: %s>" % src)
            else:
                # avoid exposing filepaths in document
                self.doc.add_paragraph("<image: %s>" % get_filename_from_url(src))
        # add styles?

    def handle_table(self):
        """
        To handle nested tables, we will parse tables manually as follows:
        Get table soup
        Create docx table
        Iterate over soup and fill docx table with new instances of this parser
        Tell HTMLParser to ignore any tags until the corresponding closing table tag
        """
        table_soup = self.tables[self.table_no]
        rows, cols = self.get_table_dimensions(table_soup)
        self.table = self.doc.add_table(rows, cols)
        rows = self.get_table_rows(table_soup)
        cell_row = 0
        for row in rows:
            cols = self.get_table_columns(row)
            cell_col = 0
            for col in cols:
                cell_html = self.get_cell_html(col)
                if col.name == 'th':
                    cell_html = "<b>%s</b>" % cell_html
                docx_cell = self.table.cell(cell_row, cell_col)
                child_parser = HtmlToDocx()
                child_parser.add_html_to_cell(cell_html, docx_cell)
                cell_col += 1
            cell_row += 1
        
        # skip all tags until corresponding closing tag
        self.instances_to_skip = len(table_soup.find_all('table'))
        self.skip_tag = 'table'
        self.skip = True
        self.table = None

    def handle_link(self, href, text):
        # Link requires a relationship
        is_external = href.startswith('http')
        rel_id = self.paragraph.part.relate_to(
            href,
            docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK,
            is_external=True  # don't support anchor links for this library yet
        )

        # Create the w:hyperlink tag and add needed values
        hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
        hyperlink.set(docx.oxml.shared.qn('r:id'), rel_id)


        # Create sub-run
        subrun = self.paragraph.add_run()
        rPr = docx.oxml.shared.OxmlElement('w:rPr')
        rPr.style = 'Hyperlink'
        subrun._r.append(rPr)
        subrun._r.text = text

        # Add subrun to hyperlink
        hyperlink.append(subrun._r)

        # Add hyperlink to run
        self.paragraph._p.append(hyperlink)

    def handle_starttag(self, tag, attrs):
        if self.skip:
            return
        if tag == 'head':
            self.skip = True
            self.skip_tag = tag
            self.instances_to_skip = 0
            return
        elif tag == 'body':
            return

        current_attrs = dict(attrs)

        if tag == 'span':
            self.tags['span'].append(current_attrs)
            return
        elif tag == 'ol' or tag == 'ul':
            self.tags['list'].append(tag)
            return # don't apply styles for now
        elif tag == 'br':
            self.run.add_break()
            return

        self.tags[tag] = current_attrs
        if tag == 'p':
            self.paragraph = self.doc.add_paragraph()

        elif tag == 'li':
            self.handle_li()

        elif tag[0] == 'h' and len(tag) == 2:
            if isinstance(self.doc, docx.document.Document):
                h_size = int(tag[1])
                self.paragraph = self.doc.add_heading(level=min(h_size, 9))
            else:
                self.paragraph = self.doc.add_paragraph()

        elif tag == 'img':
            self.handle_img(current_attrs)
            return

        elif tag == 'table':
            self.handle_table()
            return

        # set new run reference point in case of leading line breaks
        if tag == 'p' or tag == 'li':
            self.run = self.paragraph.add_run()

        # add style
        if not self.include_styles:
            return
        if 'style' in current_attrs and self.paragraph:
            style = self.parse_dict_string(current_attrs['style'])
            self.add_styles_to_paragraph(style)

    def handle_endtag(self, tag):
        if self.skip:
            if not tag == self.skip_tag:
                return

            if self.instances_to_skip > 0:
                self.instances_to_skip -= 1
                return

            self.skip = False
            self.skip_tag = None
            self.paragraph = None

        if tag == 'span':
            if self.tags['span']:
                self.tags['span'].pop()
                return
        elif tag == 'ol' or tag == 'ul':
            remove_last_occurence(self.tags['list'], tag)
            return
        elif tag == 'table':
            self.table_no += 1
            self.table = None
            self.doc = self.document
            self.paragraph = None

        if tag in self.tags:
            self.tags.pop(tag)
        # maybe set relevant reference to None?

    def handle_data(self, data):
        if self.skip:
            return

        if not self.paragraph:
            self.paragraph = self.doc.add_paragraph()

        # There can only be one nested link in a valid html document
        # You cannot have interactive content in an A tag, this includes links
        # https://html.spec.whatwg.org/#interactive-content
        link = self.tags.get('a')
        if link:
            self.handle_link(link['href'], data)
        else:
            # If there's a link, dont put the data directly in the run
            self.run = self.paragraph.add_run(data)
            spans = self.tags['span']
            for span in spans:
                if 'style' in span:
                    style = self.parse_dict_string(span['style'])
                    self.add_styles_to_run(style)


            # add font style
            for tag in self.tags:
                if tag in fonts:
                    font_style = fonts[tag]
                    setattr(self.run.font, font_style, True)

    def ignore_nested_tables(self, tables_soup):
        """
        Returns array containing only the highest level tables
        Operates on the assumption that bs4 returns child elements immediately after
        the parent element in `find_all`. If this changes in the future, this method will need to be updated

        :return:
        """
        new_tables = []
        nest = 0
        for table in tables_soup:
            if nest:
                nest -= 1
                continue
            new_tables.append(table)
            nest = len(table.find_all('table'))
        return new_tables

    def get_table_rows(self, table_soup):
        # If there's a header, body, footer or direct child tr tags, add row dimensions from there
        return table_soup.select(', '.join(self.table_row_selectors), recursive=False)

    def get_table_columns(self, row):
        # Get all columns for the specified row tag.
        return row.find_all(['th', 'td'], recursive=False) if row else []

    def get_table_dimensions(self, table_soup):
        # Get rows for the table
        rows = self.get_table_rows(table_soup)
        # Table is either empty or has non-direct children between table and tr tags
        # Thus the row dimensions and column dimensions are assumed to be 0

        cols = self.get_table_columns(rows[0]) if rows else []
        return len(rows), len(cols)

    def get_tables(self):
        if not hasattr(self, 'soup'):
            self.include_tables = False
            return
            # find other way to do it, or require this dependency?
        self.tables = self.ignore_nested_tables(self.soup.find_all('table'))  
        self.table_no = 0

    def run_process(self, html):
        if self.bs and BeautifulSoup:
            self.soup = BeautifulSoup(html, 'html.parser')
            html = remove_whitespace(str(self.soup))
        else:
            html = remove_whitespace(html)
        if self.include_tables:
            self.get_tables()
        self.feed(html)

    def add_html_to_document(self, html, document):
        if not isinstance(html, str):
            raise ValueError('First argument needs to be a %s' % str)
        elif not isinstance(document, docx.document.Document) and not isinstance(document, docx.table._Cell):
            raise ValueError('Second argument needs to be a %s' % docx.document.Document)
        self.set_initial_attrs(document)
        self.run_process(html)

    def add_html_to_cell(self, html, cell):
        if not isinstance(cell, docx.table._Cell):
            raise ValueError('Second argument needs to be a %s' % docx.table._Cell)
        unwanted_paragraph = cell.paragraphs[0]
        delete_paragraph(unwanted_paragraph)
        self.set_initial_attrs(cell)
        self.run_process(html)
        # cells must end with a paragraph or will get message about corrupt file
        # https://stackoverflow.com/a/29287121
        if not self.doc.paragraphs:
            self.doc.add_paragraph('')  

    def parse_html_file(self, filename_html, filename_docx=None):
        with open(filename_html, 'r') as infile:
            html = infile.read()
        self.set_initial_attrs()
        self.run_process(html)
        if not filename_docx:
            path, filename = os.path.split(filename_html)
            filename_docx = '%s/new_docx_file_%s' % (path, filename)
        self.doc.save('%s.docx' % filename_docx)
    
    def parse_html_string(self, html):
        self.set_initial_attrs()
        self.run_process(html)
        return self.doc

if __name__=='__main__':
    
    arg_parser = argparse.ArgumentParser(description='Convert .html file into .docx file with formatting')
    arg_parser.add_argument('filename_html', help='The .html file to be parsed')
    arg_parser.add_argument(
        'filename_docx', 
        nargs='?', 
        help='The name of the .docx file to be saved. Default new_docx_file_[filename_html]', 
        default=None
    )
    arg_parser.add_argument('--bs', action='store_true', 
        help='Attempt to fix html before parsing. Requires bs4. Default True')

    args = vars(arg_parser.parse_args())
    file_html = args.pop('filename_html')
    html_parser = HtmlToDocx()
    html_parser.parse_html_file(file_html, **args)
